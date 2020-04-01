from docx import Document
from docx.shared import Inches
import os
import requests
import io

document = Document()

document.add_heading('Titulo document',0)

p = document.add_paragraph('Un párrafo simple que tiene algo de ')
p.add_run('negritas').bold = True
p.add_run(' y algunas ')
p.add_run('italicas').italic = True

document.add_heading('Encabezado nivel 1', level=1)
document.add_paragraph('Cita intensa', style='Intense Quote')

document.add_paragraph('primer elemento en la lista desordenada',style='List Bullet')
document.add_paragraph('primer elemento en la lista ordenada',style='List Number')
image = 'https://laverdadnoticias.com/__export/1578967819952/sites/laverdad/img/2020/01/13/kyokou-suiri.png_423682103.png'
image_name = os.path.split(image)[1]
r2 = requests.get(image)
image = io.BytesIO(r2.content)
try:
    document.add_picture(image,width=Inches(1.25))
except:
    pass

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo3.docx')