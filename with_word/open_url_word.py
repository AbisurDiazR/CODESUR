import sys

from PyQt5.QtCore import QUrl
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QHBoxLayout, QLineEdit
from PyQt5.QtWidgets import QMainWindow, QPushButton, QVBoxLayout
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtWebEngineWidgets import QWebEngineView
# Tweepy
import tweepy
# Time
import time
# Creacion de carpeta
import os
import errno
# creacion del documento
from docx import Document
from docx.shared import Inches

consumer_key = 'gsswiM06At2InB2hgzwfpAiVO'
consumer_secret = 'jvt4RD4s6rzCUbRq4cCQWTS0dwg809TieyIUpPj2kV1UViuqbt'
access_token = '2460423055-aoTaKilqm8RCiwXWXg5d9L0Y3JF6rhVnDA5jpLl'
access_token_secret = '5IuyQNSDleh6PkS1HXSE8N1Au30JgoLhHoj9QtiI3pMhd'

auth = tweepy.OAuthHandler(consumer_key, consumer_secret)
auth.set_access_token(access_token, access_token_secret)
# con este objeto realizaremos todas las llamadas al API
api = tweepy.API(auth, wait_on_rate_limit=True,
                 wait_on_rate_limit_notify=True, compression=True)

document = Document()


class Widgets(QMainWindow):

    def __init__(self):
        QMainWindow.__init__(self)
        self.setWindowTitle("Codesur-WordDB")
        self.widget = QWidget(self)

        # Widget para el navegador
        self.webview = QWebEngineView()
        self.webview.load(QUrl("https://twitter.com/Nadiaferron19"))
        self.webview.urlChanged.connect(self.url_changed)

        # Ir hacia atras
        self.back_button = QPushButton("<")
        self.back_button.clicked.connect(self.webview.back)

        # Ir hacia adelante
        self.forward_button = QPushButton(">")
        self.forward_button.clicked.connect(self.webview.forward)

        # Actualizar la pagina
        self.refresh_button = QPushButton("Actualizar")
        self.refresh_button.clicked.connect(self.webview.reload)

        # Barra de direcciones
        self.url_text = QLineEdit()

        # Cargar la pagina actual
        self.go_button = QPushButton("Buscar")
        self.go_button.clicked.connect(self.url_set)

        # Obtener id del usuario
        self.id_button = QPushButton("Obtener Tweets")
        self.id_button.clicked.connect(self.get_id)

        # progressbar
        self.label_title = QtWidgets.QLabel("Porcentaje de tweets")
        self.progressbar = QtWidgets.QProgressBar()
        self.progressbar.setGeometry(200, 80, 250, 20)
        self.label_descripcion = QtWidgets.QLabel("Tweets obtenidos")
        self.label_tweet = QtWidgets.QLabel("#")

        self.toplayout = QHBoxLayout()
        self.toplayout.addWidget(self.back_button)
        self.toplayout.addWidget(self.forward_button)
        self.toplayout.addWidget(self.refresh_button)
        self.toplayout.addWidget(self.url_text)
        self.toplayout.addWidget(self.go_button)
        self.toplayout.addWidget(self.id_button)

        self.low_layout = QHBoxLayout()
        self.low_layout.addWidget(self.label_title)
        self.low_layout.addWidget(self.progressbar)
        self.low_layout.addWidget(self.label_descripcion)
        self.low_layout.addWidget(self.label_tweet)

        self.layout = QVBoxLayout()
        self.layout.addLayout(self.toplayout)
        self.layout.addLayout(self.low_layout)
        self.layout.addWidget(self.webview)

        self.widget.setLayout(self.layout)
        self.setCentralWidget(self.widget)

    def url_changed(self, url):
        self.url_text.setText(url.toString())
        txt = self.url_text.text()
        self.array = txt.split('/')
        print(self.array)

    def url_set(self):
        self.webview.setUrl(QUrl(self.url_text.text()))

    def get_id(self):
        time.sleep(2)
        user = api.get_user(screen_name=self.array[-1])

        path = "C:/Users/"+os.getlogin()+"/Desktop/"+user.screen_name

        QMessageBox.information(self, "Guardado iniciado",
                                "El archivo se guardara en la carpeta "+path)

        try:
            os.mkdir(path)
        except OSError as e:
            if e.errno != errno.EEXIST:
                raise

        # results = tweepy.Cursor(api.user_timeline, id=user.id)
        self.progressbar.setMaximum(3600)

        document.add_heading(user.screen_name, 0)
        document.add_heading(user.location, level=1)
        document.add_paragraph(user.description, style='Intense Quote')

        count = 0

        # file = open(path+"/"+user.screen_name+".txt", "w", encoding='utf-8')

        for tweet in tweepy.Cursor(api.user_timeline, id=user.id, limit=100).items():
            time.sleep(0.1)
            self.progressbar.setValue(count)
            # print(' Fecha: '+str(tweet.created_at) +
            #  ' Texto: '+tweet.text+' Ubicacion: '+tweet.user.location)
            if (not tweet.retweeted) and ('RT @' not in tweet.text):
                # file.write("Fecha: "+str(tweet.created_at)+" Texto: "+str(tweet.text)+" Localizacion: "+tweet.user.location+os.linesep)
                document.add_paragraph('ID: '+tweet.id_str+' Fecha: '+str(tweet.created_at) +
                                       ' Texto: '+tweet.text+' Ubicacion: '+tweet.user.location, style='List Number')
                print('Trabajando '+str(count))
                self.label_tweet.setText(str(count))
                count += 1

        document.save(path+'/'+user.name+'.docx')        

        QMessageBox.information(
            self, "Guardado completado", "Archivo guardado en "+path)

        # file.close()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = Widgets()
    window.show()
    sys.exit(app.exec_())
